import logging
import re
from typing import Dict, List, Set

import numpy as np
import tiktoken
from scipy import spatial

from .tree_structures import Node
from langchain.document_loaders.unstructured import UnstructuredFileLoader
from typing import List
import tqdm
from unstructured.partition.text import partition_text

logging.basicConfig(
    format="%(asctime)s - %(pathname)s - %(message)s",
    level=logging.INFO
)


def reverse_mapping(layer_to_nodes: Dict[int, List[Node]]) -> Dict[Node, int]:
    node_to_layer = {}
    for layer, nodes in layer_to_nodes.items():
        for node in nodes:
            node_to_layer[node.index] = layer
    return node_to_layer


def split_text(
    text: str, tokenizer: tiktoken.get_encoding("cl100k_base"), max_tokens: int, overlap: int = 0
):
    """
    Splits the input text into smaller chunks based on the tokenizer and maximum allowed tokens.
    
    Args:
        text (str): The text to be split.
        tokenizer (CustomTokenizer): The tokenizer to be used for splitting the text.
        max_tokens (int): The maximum allowed tokens.
        overlap (int, optional): The number of overlapping tokens between chunks. Defaults to 0.
    
    Returns:
        List[str]: A list of text chunks.
    """
    # Split the text into sentences using multiple delimiters
    delimiters = [".", "!", "?", "\n"]
    regex_pattern = "|".join(map(re.escape, delimiters))
    sentences = re.split(regex_pattern, text)
    
    # Calculate the number of tokens for each sentence
    n_tokens = [len(tokenizer.encode(" " + sentence)) for sentence in sentences]
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    for sentence, token_count in zip(sentences, n_tokens):
        # If the sentence is empty or consists only of whitespace, skip it
        if not sentence.strip():
            continue
        
        # If the sentence is too long, split it into smaller parts
        if token_count > max_tokens:
            sub_sentences = re.split(r"[,;:]", sentence)
            sub_token_counts = [len(tokenizer.encode(" " + sub_sentence)) for sub_sentence in sub_sentences]
            
            sub_chunk = []
            sub_length = 0
            
            for sub_sentence, sub_token_count in zip(sub_sentences, sub_token_counts):
                if sub_length + sub_token_count > max_tokens:
                    chunks.append(" ".join(sub_chunk))
                    sub_chunk = sub_chunk[-overlap:] if overlap > 0 else []
                    sub_length = sum(sub_token_counts[max(0, len(sub_chunk) - overlap):len(sub_chunk)])
                
                sub_chunk.append(sub_sentence)
                sub_length += sub_token_count
            
            if sub_chunk:
                chunks.append(" ".join(sub_chunk))
        
        # If adding the sentence to the current chunk exceeds the max tokens, start a new chunk
        elif current_length + token_count > max_tokens:
            chunks.append(" ".join(current_chunk))
            current_chunk = current_chunk[-overlap:] if overlap > 0 else []
            current_length = sum(n_tokens[max(0, len(current_chunk) - overlap):len(current_chunk)])
            current_chunk.append(sentence)
            current_length += token_count
        
        # Otherwise, add the sentence to the current chunk
        else:
            current_chunk.append(sentence)
            current_length += token_count
    
    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    
    return chunks


def distances_from_embeddings(
    query_embedding: List[float],
    embeddings: List[List[float]],
    distance_metric: str = "cosine",
) -> List[float]:
    """
    Calculates the distances between a query embedding and a list of embeddings.

    Args:
        query_embedding (List[float]): The query embedding.
        embeddings (List[List[float]]): A list of embeddings to compare against the query embedding.
        distance_metric (str, optional): The distance metric to use for calculation. Defaults to 'cosine'.

    Returns:
        List[float]: The calculated distances between the query embedding and the list of embeddings.
    """
    distance_metrics = {
        "cosine": spatial.distance.cosine,
        "L1": spatial.distance.cityblock,
        "L2": spatial.distance.euclidean,
        "Linf": spatial.distance.chebyshev,
    }

    if distance_metric not in distance_metrics:
        raise ValueError(
            f"Unsupported distance metric '{distance_metric}'. Supported metrics are: {list(distance_metrics.keys())}"
        )

    distances = [
        distance_metrics[distance_metric](query_embedding, embedding)
        for embedding in embeddings
    ]

    return distances


def get_node_list(node_dict: Dict[int, Node]) -> List[Node]:
    """
    Converts a dictionary of node indices to a sorted list of nodes.

    Args:
        node_dict (Dict[int, Node]): Dictionary of node indices to nodes.

    Returns:
        List[Node]: Sorted list of nodes.
    """
    indices = sorted(node_dict.keys())
    node_list = [node_dict[index] for index in indices]
    return node_list


def get_embeddings(node_list: List[Node], embedding_model: str) -> List:
    """
    Extracts the embeddings of nodes from a list of nodes.

    Args:
        node_list (List[Node]): List of nodes.
        embedding_model (str): The name of the embedding model to be used.

    Returns:
        List: List of node embeddings.
    """
    return [node.embeddings[embedding_model] for node in node_list]


def get_children(node_list: List[Node]) -> List[Set[int]]:
    """
    Extracts the children of nodes from a list of nodes.

    Args:
        node_list (List[Node]): List of nodes.

    Returns:
        List[Set[int]]: List of sets of node children indices.
    """
    return [node.children for node in node_list]


def get_text(node_list: List[Node]) -> str:
    """
    Generates a single text string by concatenating the text from a list of nodes.

    Args:
        node_list (List[Node]): List of nodes.

    Returns:
        str: Concatenated text.
    """
    text = ""
    for node in node_list:
        text += f"{' '.join(node.text.splitlines())}"
        text += "\n\n"
    return text


def indices_of_nearest_neighbors_from_distances(distances: List[float]) -> np.ndarray:
    """
    Returns the indices of nearest neighbors sorted in ascending order of distance.

    Args:
        distances (List[float]): A list of distances between embeddings.

    Returns:
        np.ndarray: An array of indices sorted by ascending distance.
    """
    return np.argsort(distances)

class RapidOCRDocLoader(UnstructuredFileLoader):
    """
    rapid ocr doc loader
    """
    def _get_elements(self) -> List:
        """
        get elements
        """
        def doc2text(filepath):
            """
            doc to text
            """
            from docx.table import _Cell, Table
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.text.paragraph import Paragraph
            from docx import Document, ImagePart
            from PIL import Image
            from io import BytesIO
            import numpy as np
            from rapidocr_onnxruntime import RapidOCR
            ocr = RapidOCR()
            doc = Document(filepath)
            resp = ""

            def iter_block_items(parent):
                """
                iter block items
                """
                from docx.document import Document
                if isinstance(parent, Document):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    parent_elm = parent._tc
                else:
                    raise ValueError("RapidOCRDocLoader parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            b_unit = tqdm.tqdm(total=len(doc.paragraphs) + len(doc.tables),
                               desc="RapidOCRDocLoader block index: 0")
            for i, block in enumerate(iter_block_items(doc)):
                b_unit.set_description(
                    "RapidOCRDocLoader  block index: {}".format(i))
                b_unit.refresh()
                if isinstance(block, Paragraph):
                    resp += block.text.strip() + "\n"
                    images = block._element.xpath('.//pic:pic')  # 获取所有图片
                    for image in images:
                        for img_id in image.xpath('.//a:blip/@r:embed'):  # 获取图片id
                            part = doc.part.related_parts[img_id]  # 根据图片id获取对应的图片
                            if isinstance(part, ImagePart):
                                image = Image.open(BytesIO(part._blob))
                                result, _ = ocr(np.array(image))
                                if result:
                                    ocr_result = [line[1] for line in result]
                                    resp += "\n".join(ocr_result)
                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                resp += paragraph.text.strip() + "\n"
                b_unit.update(1)
            return resp

        text = doc2text(self.file_path)
        return partition_text(text=text, **self.unstructured_kwargs)